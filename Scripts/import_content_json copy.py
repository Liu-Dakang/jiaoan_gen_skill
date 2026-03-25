import json
from docx import Document

def insert_row_before(table, row_idx):
    """
    核心辅助函数：在表格的指定行之前插入一个新行。
    这样可以确保新插入的教学环节行永远在“板书”和“作业”行的上方。
    """
    ref_row = table.rows[row_idx]
    new_row = table.add_row()
    # 使用 python-docx 的底层 XML 操作将新行移动到参照行之前
    ref_row._tr.addprevious(new_row._tr)
    return new_row

def fill_lesson_plan(json_path, template_path, output_path):
    # 1. 加载 JSON 数据
    with open(json_path, 'r', encoding='utf-8') as f:
        data = json.load(f)
    
    # 2. 加载 Word 模板
    doc = Document(template_path)
    
    # --- 填充 Table 0 (基础信息表) ---
    info_table = doc.tables[0]
    info_table.cell(1, 1).text = data.get("课题", "")
    info_table.cell(3, 3).text = data.get("授课日期", "")
    info_table.cell(4, 1).text = data.get("教学内容分析", "")
    info_table.cell(6, 1).text = "\n".join(data.get("学习目标", []))
    info_table.cell(7, 1).text = data.get("重点及其突破方法", "")
    info_table.cell(8, 1).text = data.get("难点及其化解方法", "")
    info_table.cell(10, 1).text = data.get("教学策略", "")
    info_table.cell(12, 1).text = data.get("学业评价", "")

    # --- 填充 Table 1 (教学过程设计表) ---
    process_table = doc.tables[1]
    process_steps = data.get("教学过程设计表格", [])
    
    # 1. 寻找“板书”所在的初始行索引
    board_row_idx = -1
    for i, row in enumerate(process_table.rows):
        if "板书" in row.cells[0].text:
            board_row_idx = i
            break
            
    # 从第2行（索引为2）开始填充
    current_row_idx = 2
    
    # 2. 填充数据，必要时插入新行
    for step in process_steps:
        if current_row_idx >= board_row_idx:
            insert_row_before(process_table, board_row_idx)
            board_row_idx += 1
        
        row_cells = process_table.rows[current_row_idx].cells
        row_cells[0].text = step.get("教学环节", "")
        row_cells[1].text = step.get("学习内容", "")
        row_cells[2].text = step.get("学生活动", "")
        row_cells[3].text = step.get("教师活动", "")
        row_cells[4].text = step.get("设计意图", "")
        
        current_row_idx += 1

    # --- 核心修改：删除板书行以上多余的空行 ---
    # 如果数据填完了，但 current_row_idx 还没到 board_row_idx，说明中间有空行
    while current_row_idx < board_row_idx:
        # 始终删除当前填充位置到板书行之间的那一行
        row_to_remove = process_table.rows[current_row_idx]
        row_to_remove._element.getparent().remove(row_to_remove._element)
        # 删除一行后，板书行的索引会减 1
        board_row_idx -= 1
        
    # --- 填充“板书”与“作业”单元格 ---
    # 由于插入了新行，索引可能发生了变化，这里直接遍历定位到单元格进行写入
    for row in process_table.rows:
        if "板书" in row.cells[0].text:
            # 填入板书对应的第二列
            row.cells[1].text = data.get("板书", "")
        elif "作业" in row.cells[0].text:
            # 填入作业对应的第二列
            row.cells[1].text = data.get("作业", "")

    # 3. 保存文件
    doc.save(output_path)
    print(f"教案已成功生成，动态行插入完成：{output_path}")
    from docx2pdf import convert
    # 单文件转换
    convert(output_path, "output.pdf")

if __name__ == "__main__":
    base_dir = r"C:\Users\22903\Desktop\CamScanner\.opencode\skills\.agents\skills\jiaoan_gen_skill"
    fill_lesson_plan(
        f'{base_dir}\\assets\\data.json', 
        f'{base_dir}\\references\\云服务实践-教案模板.docx', 
        f'{base_dir}\\教案_带作业板书输出.docx'
    )