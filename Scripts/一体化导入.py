import json
from docx import Document

def fill_template_bulletproof(json_path, template_path, output_path):
    print("正在加载 JSON 数据...")
    with open(json_path, 'r', encoding='utf-8') as f:
        data = json.load(f)

    print("正在加载 Word 模板...")
    doc = Document(template_path)

    # 1. 精确的替换字典匹配
    replace_dict = {
        "{课程名称}": data.get("{课程名称}", ""),
        "{学习任务}": data.get("{学习任务}", ""),
        "{总课时}": data.get("{总课时}", ""),
        "{学习活动}": data.get("{学习活动}", ""),
        "{分课时}": data.get("{分课时}", ""),
        "{授课日期}": data.get("{授课日期}", ""),
        "{学习任务描述}": data.get("{学习任务描述}", ""),
        "{本次学习活动}": data.get("{本次学习活动}", ""),
        "{本次学习目标}": data.get("{本次学习目标}", ""),
        "{本次学习内容}": data.get("{本次学习内容}", ""),
        "{重点内容}": data.get("{重点内容}", ""),
        "{突破方法}": data.get("{突破方法}", ""),
        "{难点内容}": data.get("{难点内容}", ""),
        "{化解方法}": data.get("{化解方法}", ""),
        "{教学策略}": data.get("{教学策略}", ""),
        "{阶段性学业成果}": data.get("{阶段性学业成果}", ""),
        "{教学过程设计}": data.get("{教学过程设计}", ""), 
        "{学习效果评价}": data.get("{学习效果评价}", ""),
        "{板书}": data.get("{板书}", ""),
        "{作业}": data.get("{作业}", "")
    }

    print("正在替换基础信息 (已启用合并单元格保护机制)...")
    basis_counter = 0 
    
    # 2. 遍历全文，使用 processed_cells 集合防止合并单元格重复写入
    processed_cells = set()
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                # 获取底层 XML 对象的内存地址，识别是否为同一个合并单元格
                if cell._tc in processed_cells:
                    continue
                processed_cells.add(cell._tc)
                
                # 对单元格内的段落进行文字替换，最大限度保护原有字体格式
                for paragraph in cell.paragraphs:
                    # 处理 {确定依据} 顺序问题
                    if "{确定依据}" in paragraph.text:
                        if basis_counter == 0:
                            paragraph.text = paragraph.text.replace("{确定依据}", data.get("key_points_basis", ""))
                            basis_counter += 1
                        else:
                            paragraph.text = paragraph.text.replace("{确定依据}", data.get("diff_points_basis", ""))
                            
                    # 替换其他标签
                    for marker, value in replace_dict.items():
                        if marker in paragraph.text:
                            paragraph.text = paragraph.text.replace(marker, str(value))

    # 3. 填写教学过程设计表格
    print("正在填充教学过程列表...")
    process_table = None
    for table in doc.tables:
        try:
            if "教学环节" in table.cell(1, 0).text:
                process_table = table
                break
        except Exception:
            continue
            
    if process_table:
        process_data = data.get("teaching_process", [])
        
        board_row_idx = -1
        for i, row in enumerate(process_table.rows):
            if "板书" in row.cells[0].text:
                board_row_idx = i
                break
        
        if board_row_idx != -1:
            current_row = 2 
            for step in process_data:
                if current_row < board_row_idx:
                    process_table.cell(current_row, 0).text = step.get("step", "")
                    process_table.cell(current_row, 1).text = step.get("content", "")
                    process_table.cell(current_row, 2).text = step.get("student_act", "")
                    process_table.cell(current_row, 3).text = step.get("teacher_act", "")
                    process_table.cell(current_row, 4).text = step.get("method_intent", "")
                    current_row += 1
                else:
                    print(f"⚠️ 警告：模板中预留的空白行用完了，无法写入: {step.get('step')}")
            
            # 删除多余空白行
            if current_row < board_row_idx:
                rows_to_delete = process_table.rows[current_row:board_row_idx]
                deleted_count = 0
                for row in rows_to_delete:
                    try:
                        tbl = row._element.getparent()
                        tbl.remove(row._element)
                        deleted_count += 1
                    except Exception:
                        pass
                print(f"🧹 已清理 {deleted_count} 行多余空白行。")
    else:
        print("⚠️ 警告：未能自动找到包含【教学环节】的表格，跳过该部分写入。")

    print("正在保存文件...")
    try:
        doc.save(output_path)
        from docx2pdf import convert
    # 单文件转换
        convert(output_path, "output.pdf")
        print(f"✅ 成功！教案已生成并保存为: {output_path}")
    except PermissionError:
        print(f"❌ [保存失败] 文件正在被使用，请关闭已打开的 Word 窗口后再运行程序！")
    except Exception as e:
        print(f"❌ [保存失败] 发生未知错误: {e}")
        doc.save(output_path)


if __name__ == "__main__":
    JSON_FILE = "assets\\一体化模板.json"           
    TEMPLATE_FILE = "references\\一体化教案空白模板.docx" 
    OUTPUT_FILE = "新生成_工学一体化教案.docx" 
    
    fill_template_bulletproof(JSON_FILE, TEMPLATE_FILE, OUTPUT_FILE)